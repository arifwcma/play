from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document


def generate(output_path: Path) -> None:
    document = Document()

    today = date.today().strftime("%d %B %Y")

    document.add_paragraph(today)
    document.add_paragraph("")

    document.add_paragraph("Apple Developer Program Support")
    document.add_paragraph("Re: Apple Developer Program enrolment letter of support")
    document.add_paragraph("Enrolment ID: 89YH793A35")
    document.add_paragraph("")

    document.add_paragraph("To whom it may concern,")
    document.add_paragraph("")

    document.add_paragraph(
        "This letter confirms that Mohammad Arifur Rahman has the legal authority to bind "
        "Wimmera Catchment Management Authority to all legal agreements presented on behalf of "
        "the Apple Developer Program."
    )
    document.add_paragraph(
        "Wimmera Catchment Management Authority intends to participate in the Apple Developer Program."
    )
    document.add_paragraph("")

    document.add_paragraph("Sincerely,")
    document.add_paragraph("")
    document.add_paragraph("[Signature]")
    document.add_paragraph("Nicole Netherway")
    document.add_paragraph("Business Manager")
    document.add_paragraph("Wimmera Catchment Management Authority")
    document.add_paragraph("Nicole.Netherway@wcma.vic.gov.au")
    document.add_paragraph("+61353829941")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))


if __name__ == "__main__":
    generate(Path("apple-developer-letter-of-support.docx"))
